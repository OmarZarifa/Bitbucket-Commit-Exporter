from docx import Document
import os

def write_commits_to_doc(filename, commits):
    """
    Creates (or appends to) a .docx and writes each commit as a paragraph.
    If the file exists, it will append; otherwise it will create a new doc.
    """
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
        doc.add_heading('Bitbucket Commits', level=1)

    for commit in commits:
        line = f"{commit['hash']} | {commit['author']} | {commit['date']} | {commit['message']}"
        doc.add_paragraph(line)

    doc.save(filename)
    print(f"Wrote {len(commits)} commits to {filename}")
